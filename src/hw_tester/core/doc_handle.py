"""
Doc Handle - Word report generation for HW Tester.
"""
from __future__ import annotations

from pathlib import Path
from copy import deepcopy
from typing import Dict, List, Optional, Tuple
from datetime import datetime

import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PySide6.QtWidgets import QApplication, QFileDialog


DEFAULT_OUTPUT_DIR = Path(r"C:\ArduinoProject\IO_Tester\tests\Results\Base_Doc\template_docs")
DEFAULT_GENERAL_DATA_PATH = DEFAULT_OUTPUT_DIR / "general_data.docx"
DEFAULT_GENERAL_OPEN_DATA_PATH = DEFAULT_OUTPUT_DIR / "open_general_data.docx"
CONFIGURATION = DEFAULT_OUTPUT_DIR / "configuration.jpg"


def _get_or_create_qt_app() -> tuple[QApplication, bool]:
	"""Return QApplication instance and whether it was created here."""
	app = QApplication.instance()
	if app is None:
		app = QApplication([])
		return app, True
	return app, False


def _parse_filename_details(file_path: Path) -> Tuple[str, str]:
	"""Extract connector name and test date/time string from file name."""
	stem_parts = file_path.stem.split("_")
	if len(stem_parts) >= 3:
		date_part = stem_parts[-2]
		time_part = stem_parts[-1]
		if date_part.isdigit() and time_part.isdigit() and len(date_part) == 8 and len(time_part) == 6:
			connector_name = "_".join(stem_parts[:-2])
			try:
				dt = datetime.strptime(f"{date_part}_{time_part}", "%Y%m%d_%H%M%S")
				return connector_name, dt.strftime("%Y-%m-%d %H:%M:%S")
			except ValueError:
				pass

	return file_path.stem, "Unknown"


def _load_report_rows(file_path: Path) -> Tuple[List[str], List[Dict[str, str]]]:
	"""Load report rows from Excel file."""
	wb = openpyxl.load_workbook(file_path, data_only=True)
	ws = wb.active

	rows = list(ws.iter_rows(values_only=True))
	wb.close()

	if not rows:
		return [], []

	headers = [str(h).strip() if h is not None else "" for h in rows[0]]
	headers = [h for h in headers if h]
	if not headers:
		return [], []

	data_rows: List[Dict[str, str]] = []
	for row in rows[1:]:
		row_dict: Dict[str, str] = {}
		for idx, header in enumerate(headers):
			value = row[idx] if idx < len(row) else ""
			row_dict[header] = "" if value is None else str(value)
		data_rows.append(row_dict)

	return headers, data_rows


def _analyze_report(rows: List[Dict[str, str]]) -> Tuple[str, Dict[str, Tuple[int, int]], int]:
	"""Return summary result, pass/total counts by test type, and not-tested count."""
	counts = {
		"power": [0, 0],
		"pullup": [0, 0],
		"logic": [0, 0],
	}
	not_tested = 0
	any_fail = False

	for row in rows:
		power_result = row.get("Power_Result", "")
		pullup_result = row.get("PullUp_Result", "")
		logic_result = row.get("Logic_DI_Result", "")

		results = [power_result, pullup_result, logic_result]
		has_pass_fail = any(r in ("Pass", "Fail") for r in results)

		if not has_pass_fail:
			not_tested += 1

		if power_result in ("Pass", "Fail"):
			counts["power"][1] += 1
			if power_result == "Pass":
				counts["power"][0] += 1
			else:
				any_fail = True

		if pullup_result in ("Pass", "Fail"):
			counts["pullup"][1] += 1
			if pullup_result == "Pass":
				counts["pullup"][0] += 1
			else:
				any_fail = True

		if logic_result in ("Pass", "Fail"):
			counts["logic"][1] += 1
			if logic_result == "Pass":
				counts["logic"][0] += 1
			else:
				any_fail = True

	if any_fail:
		summary = "FAIL"
	elif not_tested > 0:
		summary = "PARTIAL"
	else:
		summary = "PASS"

	return summary, {
		"power": (counts["power"][0], counts["power"][1]),
		"pullup": (counts["pullup"][0], counts["pullup"][1]),
		"logic": (counts["logic"][0], counts["logic"][1]),
	}, not_tested


def _add_field_run(paragraph, field_name: str):
	"""Add a Word field code (PAGE, NUMPAGES) to a paragraph."""
	run = paragraph.add_run()
	field_begin = OxmlElement("w:fldChar")
	field_begin.set(qn("w:fldCharType"), "begin")
	field_instr = OxmlElement("w:instrText")
	field_instr.set(qn("xml:space"), "preserve")
	field_instr.text = field_name
	field_sep = OxmlElement("w:fldChar")
	field_sep.set(qn("w:fldCharType"), "separate")
	field_end = OxmlElement("w:fldChar")
	field_end.set(qn("w:fldCharType"), "end")
	run._r.extend([field_begin, field_instr, field_sep, field_end])
	return run


def _set_run_style(run, size_pt: int, rgb: tuple[int, int, int]) -> None:
	"""Apply font size and color to a run."""
	run.font.size = Pt(size_pt)
	run.font.color.rgb = RGBColor(*rgb)


def _add_header_footer(doc: Document) -> None:
	"""Apply header/footer layout: left title, right date, right page numbering."""
	for section in doc.sections:
		header = section.header
		footer = section.footer

		# Clear default empty paragraph text.
		if header.paragraphs:
			header.paragraphs[0].text = ""
		if footer.paragraphs:
			footer.paragraphs[0].text = ""

		usable_width = section.page_width - section.left_margin - section.right_margin

		# Header table with two columns: title left, date right.
		header_table = header.add_table(rows=1, cols=2, width=usable_width)
		header_table.autofit = True
		left_cell, right_cell = header_table.rows[0].cells

		title_run = left_cell.paragraphs[0].add_run("IO Tester - Test Report")
		_set_run_style(title_run, 12, (102, 153, 204))
		left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

		date_text = datetime.now().strftime("%Y-%m-%d")
		date_run = right_cell.paragraphs[0].add_run(date_text)
		_set_run_style(date_run, 10, (0, 0, 0))
		right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

		# Footer table: right aligned page numbering.
		footer_table = footer.add_table(rows=1, cols=1, width=usable_width)
		footer_table.autofit = True
		footer_paragraph = footer_table.rows[0].cells[0].paragraphs[0]
		footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		label_run = footer_paragraph.add_run("Page ")
		_set_run_style(label_run, 9, (128, 128, 128))
		page_run = _add_field_run(footer_paragraph, "PAGE")
		_set_run_style(page_run, 9, (128, 128, 128))
		mid_run = footer_paragraph.add_run(" of ")
		_set_run_style(mid_run, 9, (128, 128, 128))
		numpages_run = _add_field_run(footer_paragraph, "NUMPAGES")
		_set_run_style(numpages_run, 9, (128, 128, 128))


def _append_document_body(target: Document, source: Document) -> None:
	"""Append all body elements (paragraphs, tables) from source into target."""
	element_count = 0
	
	# Get all existing content before any section properties
	target_body = target.element.body
	
	# Find insertion point (before section properties if they exist)
	insert_index = len(target_body)
	for i, element in enumerate(target_body):
		if element.tag == qn("w:sectPr"):
			insert_index = i
			break
	
	# Insert source elements at the correct position
	for element in source.element.body:
		if element.tag == qn("w:sectPr"):
			continue
		target_body.insert(insert_index, deepcopy(element))
		insert_index += 1
		element_count += 1
	
	print(f"  → Appended {element_count} elements from source document")


def _add_section_heading(doc: Document, text: str) -> None:
	"""Add a blue section heading."""
	heading = doc.add_paragraph(text)
	heading.runs[0].font.size = Pt(14)
	heading.runs[0].font.color.rgb = RGBColor(68, 114, 196)  # Blue color
	heading.runs[0].bold = True
	heading.space_after = Pt(6)


# def _add_title(doc: Document, text: str) -> None:
# 	"""Add a large blue title."""
# 	title = doc.add_paragraph(text)
# 	title.runs[0].font.size = Pt(18)
# 	title.runs[0].font.color.rgb = RGBColor(68, 114, 196)  # Blue color
# 	title.runs[0].bold = True
# 	title.space_after = Pt(12)


def create_doc_report(report_files: List[Path], output_path: Path) -> Path:
	"""Create a Word document report from selected Excel files."""
	document = Document()
	_add_header_footer(document)
	
	print(f"Step 1: Creating document with header/footer")
	
	# 1. Open general data (if exists) - FIRST
	if DEFAULT_GENERAL_OPEN_DATA_PATH.exists():
		print(f"Step 2: Adding open_general_data.docx")
		general_open_doc = Document(DEFAULT_GENERAL_OPEN_DATA_PATH)
		_append_document_body(document, general_open_doc)
	else:
		print(f"Step 2: Skipping open_general_data.docx (file not found)")
	
	# 2. Configuration image (if exists) - SECOND
	if CONFIGURATION.exists():
		print(f"Step 3: Adding configuration.jpg")
		document.add_picture(str(CONFIGURATION), width=Inches(6.0))
		document.add_paragraph()  # Spacing after image
	else:
		print(f"Step 3: Skipping configuration.jpg (file not found)")
	
	# 3. Result section - Summary table - THIRD
	print(f"Step 4: Adding Result section and summary table")
	_add_section_heading(document, "Result")


	summary_table = document.add_table(rows=1, cols=8)
	summary_table.style = "Table Grid"
	summary_table.autofit = False
	summary_col_widths = [
		Inches(0.35),
		Inches(0.9),
		Inches(1.7),
		Inches(0.9),
		Inches(0.7),
		Inches(0.7),
		Inches(0.7),
		Inches(0.7),
	]

	summary_headers = [
		"#",
		"Connector",
		"Test Date/Time",
		"Summary Result",
		"Power Result",
		"PullUp Result",
		"Logic Result",
		"Not Tested",
	]
	for idx, header in enumerate(summary_headers):
		cell = summary_table.rows[0].cells[idx]
		cell.text = header
		cell.width = summary_col_widths[idx]
		# Make header bold
		for paragraph in cell.paragraphs:
			for run in paragraph.runs:
				run.font.bold = True
				run.font.size = Pt(10)

	for idx, file_path in enumerate(report_files, start=1):
		connector_name, test_datetime = _parse_filename_details(file_path)
		headers, rows = _load_report_rows(file_path)
		summary, counts, not_tested = _analyze_report(rows)

		row_cells = summary_table.add_row().cells
		row_cells[0].text = str(idx)
		row_cells[1].text = connector_name
		row_cells[2].text = test_datetime
		row_cells[3].text = summary
		row_cells[4].text = f"{counts['power'][0]}/{counts['power'][1]}"
		row_cells[5].text = f"{counts['pullup'][0]}/{counts['pullup'][1]}"
		row_cells[6].text = f"{counts['logic'][0]}/{counts['logic'][1]}"
		row_cells[7].text = str(not_tested)
		
		# Format cell content and set column widths
		for col_idx, width in enumerate(summary_col_widths):
			row_cells[col_idx].width = width
			# Set font size for all cells
			for paragraph in row_cells[col_idx].paragraphs:
				for run in paragraph.runs:
					run.font.size = Pt(9)
			
			# Color code the Summary Result column
			if col_idx == 3:  # Summary Result column
				for paragraph in row_cells[col_idx].paragraphs:
					for run in paragraph.runs:
						if summary == "PASS":
							run.font.color.rgb = RGBColor(0, 176, 80)  # Green
							run.font.bold = True
						elif summary == "FAIL":
							run.font.color.rgb = RGBColor(255, 0, 0)  # Red
							run.font.bold = True
						elif summary == "PARTIAL":
							run.font.color.rgb = RGBColor(255, 192, 0)  # Orange
							run.font.bold = True

	document.add_paragraph()  # Spacing
	
	# 4. General data / Test procedures (if exists) - FOURTH (LAST)
	if DEFAULT_GENERAL_DATA_PATH.exists():
		print(f"Step 5: Adding general_data.docx")
		general_doc = Document(DEFAULT_GENERAL_DATA_PATH)
		_append_document_body(document, general_doc)
	else:
		print(f"Step 5: Skipping general_data.docx (file not found)")
	
	print(f"Step 6: Saving document to {output_path}")
	output_path.parent.mkdir(parents=True, exist_ok=True)
	document.save(output_path)
	return output_path


def select_report_files(initial_dir: Optional[Path] = None) -> List[Path]:
	"""Open file dialog to select Excel report files."""
	initial = initial_dir or DEFAULT_OUTPUT_DIR
	app, created_here = _get_or_create_qt_app()
	file_paths, _ = QFileDialog.getOpenFileNames(
		None,
		"Select Report Excel Files",
		str(initial),
		"Excel files (*.xlsx)"
	)
	if created_here:
		app.quit()
	return [Path(path) for path in file_paths]


def prompt_output_path(default_dir: Optional[Path] = None) -> Optional[Path]:
	"""Prompt user for Word report output path."""
	initial = default_dir or DEFAULT_OUTPUT_DIR
	app, created_here = _get_or_create_qt_app()
	output_path, _ = QFileDialog.getSaveFileName(
		None,
		"Save Word Report",
		str(initial / "IO_Tester_Report.docx"),
		"Word Document (*.docx)"
	)
	if created_here:
		app.quit()
	return Path(output_path) if output_path else None


def create_doc_report_via_dialog() -> Optional[Path]:
	"""Select files via dialog and create the Word report."""
	report_files = select_report_files()
	if not report_files:
		return None

	output_path = prompt_output_path()
	if not output_path:
		return None

	return create_doc_report(report_files, output_path)
