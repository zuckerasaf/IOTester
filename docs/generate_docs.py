"""
Generate Word (.docx) versions of Quick_Start_Guide and User_Manual.
Run once to produce the files; delete this script after.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
C_TITLE      = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
C_HEADING    = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
C_ACCENT     = RGBColor(0x0A, 0x58, 0xCA)   # button blue
C_YELLOW_BG  = RGBColor(0xFF, 0xF5, 0x9D)   # testing yellow (for legend)
C_GREEN_BG   = RGBColor(0x90, 0xEE, 0x90)   # pass green
C_PINK_BG    = RGBColor(0xFF, 0xB6, 0xC1)   # fail pink
C_GRAY       = RGBColor(0x44, 0x44, 0x44)   # body text dark
C_TH_BG      = RGBColor(0x2E, 0x74, 0xB5)   # table header bg
C_ROW_ALT    = RGBColor(0xDE, 0xE9, 0xF7)   # table alternate row

# ─────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background with a solid colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="BFBFBF", sz="4"):
    """Apply thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, enabled in [('top', top), ('bottom', bottom),
                          ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        if enabled:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run_bold(para, text, size_pt=10, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_run_normal(para, text, size_pt=10, color=None, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    run.italic = italic
    return run


def set_paragraph_spacing(para, before=0, after=4, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(before * 20))
    pSpacing.set(qn('w:after'), str(after * 20))
    if line_spacing:
        pSpacing.set(qn('w:line'), str(int(line_spacing * 240)))
        pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)


def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_paragraph_spacing(p, before=2, after=2)
    return p


# ─────────────────────────────────────────────
# Shared style helpers
# ─────────────────────────────────────────────

def title_block(doc, title_text, subtitle_text):
    """Add the document title and subtitle."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C_TITLE
    set_paragraph_spacing(p, before=0, after=6)

    add_horizontal_rule(doc)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_normal(p2, subtitle_text, size_pt=10, color=C_GRAY, italic=True)
    set_paragraph_spacing(p2, before=0, after=12)


def section_heading(doc, number, title, level=1):
    """Add a numbered section heading."""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(f"{number}  {title}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = C_HEADING
        add_horizontal_rule(doc)
    else:
        run = p.add_run(f"{number}  {title}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = C_ACCENT
    set_paragraph_spacing(p, before=10, after=4)
    return p


def step_heading(doc, step_label, title):
    """Coloured step heading for Quick Start Guide."""
    p = doc.add_paragraph()
    run1 = p.add_run(f"{step_label}  —  ")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = C_HEADING
    run2 = p.add_run(title)
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = C_HEADING
    add_horizontal_rule(doc)
    set_paragraph_spacing(p, before=10, after=4)


def body_para(doc, text, size_pt=10):
    p = doc.add_paragraph()
    add_run_normal(p, text, size_pt=size_pt, color=C_GRAY)
    set_paragraph_spacing(p, before=0, after=4)
    return p


def bullet(doc, text, level=0, size_pt=10):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    add_run_normal(p, text, size_pt=size_pt, color=C_GRAY)
    set_paragraph_spacing(p, before=0, after=2)
    return p


def numbered_item(doc, num, text, size_pt=10):
    p = doc.add_paragraph(style='List Number')
    add_run_normal(p, text, size_pt=size_pt, color=C_GRAY)
    set_paragraph_spacing(p, before=0, after=2)
    return p


def inline_code(para, text):
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return run


def note_box(doc, text):
    """Grey-shaded note paragraph."""
    p = doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xED, 0xF2, 0xF8))
    cp = cell.paragraphs[0]
    add_run_normal(cp, text, size_pt=9, color=C_GRAY, italic=True)
    set_paragraph_spacing(cp, before=4, after=4)
    doc.add_paragraph()


def th_cell(cell, text):
    set_cell_bg(cell, C_TH_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def td_cell(cell, text, alt=False, bold=False, size_pt=9):
    if alt:
        set_cell_bg(cell, C_ROW_ALT)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = C_GRAY


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 1 — Quick Start Guide
# ─────────────────────────────────────────────────────────────────────────────

import docx

def build_quick_start():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    title_block(doc, "IO TESTER", "QUICK START GUIDE  |  Version 1.0  |  June 2026")

    # ── Placeholder image note ────────────────────────────────────────────────
    note_box(doc, "[ Insert application screenshot here ]")

    # ── Checklist ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run_bold(p, "BEFORE YOU BEGIN — CHECKLIST", size_pt=12, color=C_HEADING)
    add_horizontal_rule(doc)
    set_paragraph_spacing(p, before=6, after=4)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    widths = [Cm(1.2), Cm(13)]
    checks = [
        "Controllino Mega is powered on and connected via USB (default: COM5)",
        "IO Cards (C1–C5) are powered and connected to the test network",
        "The Unit Under Test (UUT) is connected to the test fixture",
        "The correct Excel test file (.xlsx) for this connector is available",
    ]
    for i, (row, text) in enumerate(zip(tbl.rows, checks)):
        alt = (i % 2 == 1)
        set_cell_bg(row.cells[0], C_ROW_ALT if alt else RGBColor(0xFF,0xFF,0xFF))
        set_cell_bg(row.cells[1], C_ROW_ALT if alt else RGBColor(0xFF,0xFF,0xFF))
        row.cells[0].paragraphs[0].add_run("☐").font.size = Pt(14)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        td_cell(row.cells[1], text, size_pt=10)
    doc.add_paragraph()

    # ── Steps ─────────────────────────────────────────────────────────────────
    steps = [
        ("STEP 1", "LAUNCH THE APPLICATION", [
            ("body", "Double-click IO_Tester.exe to open the application."),
            ("body", "The window opens with three areas:"),
            ("bullet", "Pin table (top — most of the screen)"),
            ("bullet", "Run Controls bar (middle)"),
            ("bullet", "Operational Log (bottom)"),
            ("body", "The application will attempt to connect to the Controllino automatically."),
            ("body", "Watch the Operational Log for:  [INFO]  Hardware initialized on COM5"),
            ("note",  "If you see a WARNING about connection, go to the Troubleshooting section."),
        ]),
        ("STEP 2", "VERIFY COMMUNICATION", [
            ("body", "Click  [ Comm Check ]  in the Test/Debug group (bottom-right panel)."),
            ("bullet", "Green confirmation in the log → Controllino responding correctly."),
            ("bullet", "ERROR in the log → cable or COM port issue (see Troubleshooting)."),
        ]),
        ("STEP 3", "LOAD A TEST FILE", [
            ("body", "Click  [ Load ]  in the Connector / File group (bottom-left panel)."),
            ("body", "Select the Excel file (.xlsx) for the connector you are testing."),
            ("body", "The pin table populates with all pins and their expected values."),
            ("body", "The [ Test ] and [ Test_All ] buttons become active."),
        ]),
        ("STEP 4", "RUN THE TESTS", [
            ("subhead", "TO TEST ALL PINS"),
            ("body",   "Click  [ Test_All ]"),
            ("bullet", "The currently-tested row turns YELLOW"),
            ("bullet", "Pass results turn GREEN  |  Fail results turn PINK"),
            ("bullet", "The status bar shows Testing Pin / Power / Pullup / Logic live"),
            ("subhead", "TO TEST SELECTED PINS ONLY"),
            ("body",   "Click one or more rows in the pin table (Ctrl+click for multiple), then click  [ Test ]"),
            ("subhead", "TO STOP EARLY"),
            ("body",   "Click  [ Stop ]  — the current pin finishes then the run halts."),
        ]),
        ("STEP 5", "READ THE RESULTS", [
            ("body", "Row colour at a glance:"),
        ]),
        ("STEP 6", "SAVE A REPORT", [
            ("body", "Click  [ Report ]"),
            ("body", "Choose a file name and location. An Excel file is saved with all pin data and test results."),
        ]),
    ]

    for step_label, step_title, items in steps:
        step_heading(doc, step_label, step_title)

        if step_label == "STEP 5":
            # Colour legend table
            colour_rows = [
                (C_YELLOW_BG, "YELLOW", "Pin is currently being tested"),
                (C_GREEN_BG,  "GREEN",  "At least one test passed, none failed"),
                (C_PINK_BG,   "PINK",   "At least one test failed"),
                (RGBColor(0x2A,0x2A,0x2A), "DARK", "Pin has not been tested yet"),
            ]
            legend = doc.add_table(rows=len(colour_rows), cols=2)
            legend.style = 'Table Grid'
            for row_data, row in zip(colour_rows, legend.rows):
                bg, label, desc = row_data
                set_cell_bg(row.cells[0], bg)
                p0 = row.cells[0].paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r0 = p0.add_run(label)
                r0.bold = True
                r0.font.size = Pt(9)
                if label == "DARK":
                    r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                td_cell(row.cells[1], desc, size_pt=10)
            doc.add_paragraph()

            body_para(doc, "Key result columns:")
            bullet(doc, "Power_Result  /  Power_Result_Reason")
            bullet(doc, "PullUp_Result  /  PullUp_Result_Reason")
            bullet(doc, "Logic_DI_Result  /  Logic_DI_Result_Reason")
            body_para(doc, "Values are Pass or Fail.  Hover the mouse over any cell to see the full text.")
            continue

        for kind, text in items:
            if kind == "body":
                body_para(doc, text)
            elif kind == "bullet":
                bullet(doc, text)
            elif kind == "note":
                note_box(doc, text)
            elif kind == "subhead":
                p = doc.add_paragraph()
                add_run_bold(p, text, size_pt=10, color=C_ACCENT)
                set_paragraph_spacing(p, before=6, after=2)

    # ── I-BIT section ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run_bold(p, "RUNNING THE I-BIT TEST (Built-In Test)", size_pt=12, color=C_HEADING)
    add_horizontal_rule(doc)
    set_paragraph_spacing(p, before=10, after=4)

    body_para(doc, "The I-BIT test checks all 50 pins on both connector systems (A and B) for "
                   "short-circuit / open-circuit faults without requiring a test file.")

    ibit_steps = [
        "Click  [ IBIT ]  in the Test/Debug group",
        "The test runs automatically for ~100 pins; watch the Operational Log for progress",
        "The IBIT button turns GREEN (all pass) or RED (failures detected)",
        "Click  [ Stop IBIT ]  to abort early if needed",
    ]
    for i, s in enumerate(ibit_steps, 1):
        numbered_item(doc, i, s)

    doc.add_paragraph()

    # ── Troubleshooting ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run_bold(p, "TROUBLESHOOTING", size_pt=12, color=C_HEADING)
    add_horizontal_rule(doc)
    set_paragraph_spacing(p, before=10, after=4)

    ts_data = [
        ('"No communication" / hardware warning at startup',
         ['Check USB cable to Controllino is seated',
          'Verify COM port: open Settings and confirm Board.Port is correct',
          'Try Comm Check button — if it fails, restart the Controllino']),
        ('"UDP binding error" on startup',
         ['Another instance of IO Tester may already be running — close it',
          'Check that no other program is using the card network ports']),
        ('All pins reporting FAIL',
         ['Check that "Simulation: Off" is shown in Test/Debug group',
          'Confirm the correct Excel test file was loaded',
          'Verify UUT is fully seated in the fixture']),
        ('Test button stays greyed out',
         ['No file has been loaded — click Load first']),
        ('Need to run without hardware (demo / office use)',
         ['Click  [ Simulation: Off ]  to toggle to  Simulation: On',
          'Tests will run with simulated voltages — no hardware required']),
    ]

    for symptom, actions in ts_data:
        p = doc.add_paragraph()
        add_run_bold(p, symptom, size_pt=10, color=C_HEADING)
        set_paragraph_spacing(p, before=6, after=2)
        for a in actions:
            bullet(doc, a)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_normal(p, "For full reference see:  User_Manual.docx", size_pt=9,
                   color=C_GRAY, italic=True)

    doc.save(r"c:\ArduinoProject\IO_Tester\docs\Quick_Start_Guide.docx")
    print("Quick_Start_Guide.docx saved.")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT 2 — User Manual
# ─────────────────────────────────────────────────────────────────────────────

def build_user_manual():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_block(doc, "IO TESTER", "USER MANUAL  |  Version 1.0  |  June 2026")

    # ── Table of Contents (static) ────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run_bold(p, "TABLE OF CONTENTS", size_pt=12, color=C_HEADING)
    add_horizontal_rule(doc)

    toc_entries = [
        ("1",  "Overview"),
        ("2",  "System Requirements"),
        ("3",  "Application Layout"),
        ("4",  "Connector / File Group"),
        ("5",  "Run Controls Group"),
        ("6",  "Log Group"),
        ("7",  "Test / Debug Group"),
        ("8",  "The Pin Table"),
        ("9",  "Test Types Explained"),
        ("10", "I-BIT Test"),
        ("11", "Log Window"),
        ("12", "Saving a Report"),
        ("13", "Configuration (Comm_settings.yaml)"),
        ("14", "Simulation and LocalHost Modes"),
        ("15", "Debug / Step Mode"),
        ("16", "Understanding Test Results"),
        ("17", "Troubleshooting"),
        ("18", "Keyboard and Mouse Shortcuts"),
    ]

    toc_tbl = doc.add_table(rows=len(toc_entries), cols=2)
    for i, (num, title) in enumerate(toc_entries):
        alt = (i % 2 == 1)
        row = toc_tbl.rows[i]
        if alt:
            set_cell_bg(row.cells[0], C_ROW_ALT)
            set_cell_bg(row.cells[1], C_ROW_ALT)
        p0 = row.cells[0].paragraphs[0]
        p0.add_run(num).font.size = Pt(10)
        p1 = row.cells[1].paragraphs[0]
        r = p1.add_run(title)
        r.font.size = Pt(10)
        r.font.color.rgb = C_GRAY
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 1 — Overview
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "1.", "OVERVIEW")

    note_box(doc, "[ Insert application screenshot here ]")

    body_para(doc,
        "IO Tester is a hardware verification application used to test the electrical "
        "integrity of connector pins on a Unit Under Test (UUT). It drives stimulus "
        "signals through IO cards (C1–C5) and measures the resulting voltages through "
        "a Controllino Mega microcontroller, then compares measured values to the "
        "expected values defined in an Excel test file.")
    body_para(doc, "Three test types are supported per pin:")
    bullet(doc, "Power Test   — measures DC voltage on the pin")
    bullet(doc, "Pull-Up Test — activates a pull-up and measures the resulting voltage")
    bullet(doc, "Logic Test   — commands a digital output and reads a digital input")
    body_para(doc,
        "Results are colour-coded in the pin table and can be exported to an Excel report.")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 2 — System Requirements
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "2.", "SYSTEM REQUIREMENTS")

    p = doc.add_paragraph()
    add_run_bold(p, "Hardware", size_pt=10, color=C_ACCENT)
    bullet(doc, "Controllino Mega connected via USB (serial COM port)")
    bullet(doc, "IO Cards C1–C5 connected to the test network (192.168.195.x)")
    bullet(doc, "UUT connected to the test fixture")

    p = doc.add_paragraph()
    add_run_bold(p, "Software / OS", size_pt=10, color=C_ACCENT)
    bullet(doc, "Windows 10 or later (64-bit)")
    bullet(doc, "No additional installation required — IO_Tester.exe is self-contained")

    p = doc.add_paragraph()
    add_run_bold(p, "Files required at runtime (same folder as the EXE or configured path)", size_pt=10, color=C_ACCENT)
    for f in ["config\\Comm_settings.yaml — communication settings",
              "config\\pin_map.json — Controllino pin assignments",
              "config\\board_pin_config.json — board-specific pin names",
              "config\\connector_Address_A_map.xlsx — connector A mux mapping",
              "config\\connector_Address_B_map.xlsx — connector B mux mapping"]:
        p = doc.add_paragraph(style='List Bullet')
        inline_code(p, f.split(" — ")[0])
        add_run_normal(p, "  —  " + f.split(" — ")[1], size_pt=10, color=C_GRAY)
        set_paragraph_spacing(p, before=0, after=2)

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 3 — Application Layout
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "3.", "APPLICATION LAYOUT")

    note_box(doc, "[ Insert annotated screenshot of the main window here ]")

    body_para(doc, "The main window is divided into four areas:")
    bullet(doc, "Pin Table (top, occupies most of the screen) — shows all loaded connector pins and their test results")
    bullet(doc, "Control Bar (middle) — four groups: Connector/File, Run Controls, Log, Test/Debug")
    bullet(doc, "Operational Log (bottom) — real-time event feed")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 4 — Connector / File Group
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "4.", "CONNECTOR / FILE GROUP")

    note_box(doc, "[ Insert screenshot of the Connector/File panel here ]")

    buttons = [
        ("Connector field", "Displays the name of the currently loaded test file. Read-only; updated automatically when a file is loaded."),
        ("Settings",  "Opens Comm_settings.yaml in the system text editor. Changes take effect on the next application launch. See Section 13."),
        ("Load",      "Opens a file browser to select an Excel test file (.xlsx or .xlsm). After loading, the pin table populates and Test/Test_All buttons activate."),
        ("Report",    "Saves the current pin table including all test results to an Excel file. A save dialog is shown."),
        ("DOC",       "Opens the HTML flow diagram for the selected debug scenario in your browser. A local HTTP server starts in the background automatically."),
    ]

    tbl = doc.add_table(rows=len(buttons), cols=2)
    tbl.style = 'Table Grid'
    for i, (btn, desc) in enumerate(buttons):
        alt = (i % 2 == 1)
        td_cell(tbl.rows[i].cells[0], f"[ {btn} ]", alt=alt, bold=True, size_pt=10)
        td_cell(tbl.rows[i].cells[1], desc, alt=alt, size_pt=10)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 5 — Run Controls Group
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "5.", "RUN CONTROLS GROUP")

    note_box(doc, "[ Insert screenshot of the Run Controls panel here ]")

    run_buttons = [
        ("Test",     "Runs tests on the rows currently selected in the pin table. Ctrl+click to select multiple rows. Disabled until a test file is loaded."),
        ("Test_All", "Runs tests on every pin in the table, top to bottom. Disabled until a test file is loaded."),
        ("Stop",     "Stops the running test after the current pin finishes. Active only while a test is in progress."),
    ]
    tbl = doc.add_table(rows=len(run_buttons), cols=2)
    tbl.style = 'Table Grid'
    for i, (btn, desc) in enumerate(run_buttons):
        td_cell(tbl.rows[i].cells[0], f"[ {btn} ]", alt=(i%2==1), bold=True, size_pt=10)
        td_cell(tbl.rows[i].cells[1], desc, alt=(i%2==1), size_pt=10)
    doc.add_paragraph()

    body_para(doc, "Status bar (updated live during a test run):")
    status = [
        ("Testing Pin", "ID of the pin currently under test"),
        ("Power",       "Pass / Fail result of the last power measurement"),
        ("Pullup",      "Pass / Fail result of the last pull-up measurement"),
        ("Logic",       "Pass / Fail result of the last logic measurement"),
    ]
    for label, desc in status:
        p = doc.add_paragraph(style='List Bullet')
        add_run_bold(p, label + ":  ", size_pt=10, color=C_ACCENT)
        add_run_normal(p, desc, size_pt=10, color=C_GRAY)
        set_paragraph_spacing(p, before=0, after=2)

    body_para(doc, "All labels reset to '---' when a new test sequence starts.")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 6 — Log Group
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "6.", "LOG GROUP")

    body_para(doc, "The checkboxes filter which message types appear in the Operational Log:")

    log_filters = [
        ("INF", "General information messages (hardware init, file loaded, etc.)"),
        ("SUC", "Test pass events — shown in green"),
        ("WRN", "Warnings — tests failed, pins skipped — shown in yellow/orange"),
        ("ERR", "Errors — hardware faults, file errors — shown in red"),
        ("DBG", "Verbose debug messages — shown in grey. Off by default."),
    ]
    tbl = doc.add_table(rows=len(log_filters), cols=2)
    tbl.style = 'Table Grid'
    for i, (code, desc) in enumerate(log_filters):
        td_cell(tbl.rows[i].cells[0], code, alt=(i%2==1), bold=True, size_pt=10)
        td_cell(tbl.rows[i].cells[1], desc, alt=(i%2==1), size_pt=10)
    doc.add_paragraph()

    body_para(doc, "[ Clear ] — Clears all text from the Operational Log. Does not affect test results.")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 7 — Test / Debug Group
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "7.", "TEST / DEBUG GROUP")

    note_box(doc, "[ Insert screenshot of the Test/Debug panel here ]")

    debug_buttons = [
        ("Comm Check",          "Sends a ping to the Controllino and reports success or failure in the log. Use before running tests to verify hardware connection."),
        ("IBIT",                "Starts the Built-In Test (see Section 10). Tests all 50 pins on both systems A and B. Button turns GREEN (pass) or RED (fail) when complete."),
        ("Stop IBIT",           "Stops the I-BIT test after the current pin. Active only while I-BIT is running."),
        ("Simulation: On/Off",  "Toggle. ON = simulated voltages, no hardware communication. OFF = real hardware. Current state shown on button label. Setting is saved automatically."),
        ("LocalHost",           "Toggle. Switches IO card communication between real network (192.168.195.x) and local loopback (127.0.0.1). Saved automatically."),
        ("Next",                "Used in Debug/Step mode only. Advances the test to the next checkpoint when paused."),
        ("Debug: True/False",   "Toggle. Enables step-by-step debug mode. When True, the test pauses at each checkpoint and waits for Next."),
        ("HTML file dropdown",  "Selects which flow diagram HTML file is opened by the DOC button."),
    ]
    tbl = doc.add_table(rows=len(debug_buttons), cols=2)
    tbl.style = 'Table Grid'
    for i, (btn, desc) in enumerate(debug_buttons):
        td_cell(tbl.rows[i].cells[0], f"[ {btn} ]", alt=(i%2==1), bold=True, size_pt=10)
        td_cell(tbl.rows[i].cells[1], desc, alt=(i%2==1), size_pt=10)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 8 — Pin Table
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "8.", "THE PIN TABLE")

    note_box(doc, "[ Insert screenshot of the pin table with example results here ]")

    body_para(doc, "The pin table displays all pins from the loaded test file. It has 22 columns:")

    col_groups = [
        ("PIN IDENTIFICATION",
         [("ID",            "Unique pin identifier (row number in test file)"),
          ("Connect",       "Connector identifier (e.g. J1, J2)"),
          ("Discrete Name", "Functional name of the signal (e.g. LG_HOOK_UP)"),
          ("Signal Name",   "Electrical signal name"),
          ("Plug",          "Physical plug designation (e.g. 1P2)"),
          ("Type",          "Pin type (Gnd/open, Open, RELAY, etc.)"),
          ("Pin",           "Physical pin number on the connector")]),
        ("POWER TEST",
         [("Power Expected",    "Expected voltage (V) for power test"),
          ("Power Input",       "IO card command to apply voltage (e.g. C1_AO2_10)"),
          ("Power Measured",    "Voltage actually measured"),
          ("Power Result",      "Pass / Fail"),
          ("Power Reason",      "Explanation of the result")]),
        ("PULL-UP TEST",
         [("PullUp Expected",   "Expected voltage after pull-up is activated"),
          ("PullUp Input",      "IO card command for the pull-up signal"),
          ("PullUp Measured",   "Voltage measured after pull-up"),
          ("PullUp Result",     "Pass / Fail"),
          ("PullUp Reason",     "Explanation of the result")]),
        ("LOGIC TEST",
         [("Logic Pin Input",   "IO card command for the logic stimulus"),
          ("Logic Command",     "Command sent to the IO card"),
          ("Logic Expected",    "Expected digital state (High / Low)"),
          ("Logic DI Result",   "Pass / Fail"),
          ("Logic DI Reason",   "Explanation of the result")]),
    ]

    for group_name, cols in col_groups:
        p = doc.add_paragraph()
        add_run_bold(p, group_name, size_pt=10, color=C_ACCENT)
        set_paragraph_spacing(p, before=6, after=2)
        tbl = doc.add_table(rows=1 + len(cols), cols=2)
        tbl.style = 'Table Grid'
        th_cell(tbl.rows[0].cells[0], "Column")
        th_cell(tbl.rows[0].cells[1], "Description")
        for i, (col, desc) in enumerate(cols):
            td_cell(tbl.rows[i+1].cells[0], col, alt=(i%2==1), bold=True, size_pt=9)
            td_cell(tbl.rows[i+1].cells[1], desc, alt=(i%2==1), size_pt=9)
        doc.add_paragraph()

    body_para(doc, "Row colour coding:")
    colour_data = [
        (C_YELLOW_BG, "YELLOW",   "Pin is currently being tested"),
        (C_GREEN_BG,  "GREEN",    "At least one test passed; no failures"),
        (C_PINK_BG,   "PINK",     "At least one test failed"),
        (RGBColor(0x3A,0x3A,0x3A),"DARK (alternate)", "Row has not been tested yet"),
    ]
    legend = doc.add_table(rows=len(colour_data), cols=2)
    legend.style = 'Table Grid'
    for row_data, row in zip(colour_data, legend.rows):
        bg, label, desc = row_data
        set_cell_bg(row.cells[0], bg)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9)
        if label.startswith("DARK"):
            r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        td_cell(row.cells[1], desc, size_pt=10)
    doc.add_paragraph()

    body_para(doc, "Additional table features:")
    bullet(doc, "Tooltips — hover the mouse over any cell to see the full cell text")
    bullet(doc, "Sorting — click a column header to sort; click again to reverse")
    bullet(doc, "Inline editing — double-click an editable cell (Expected, Input, Command columns) to edit its value")
    bullet(doc, "Selection — click to select a row; Ctrl+click to add rows; Shift+click for a range")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 9 — Test Types
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "9.", "TEST TYPES EXPLAINED")

    body_para(doc,
        "For each pin, the tester determines which tests to run based on whether the "
        "corresponding Expected / Input columns are filled in the test file. "
        "A column left empty or set to \"-\" means that test is skipped for that pin.")

    section_heading(doc, "9.1", "Power Test", level=2)
    body_para(doc, "Purpose: Verify that the correct DC voltage is present on the pin.")
    body_para(doc, "Procedure:")
    for s in ["The mux matrix routes the pin to the measurement ADC",
              "If Power_Input is specified, the IO card applies the stimulus voltage",
              "Voltage is measured via the Controllino analog input",
              "Measured value is compared to Power_Expected ± tolerance",
              "Result: Pass if within tolerance, Fail otherwise"]:
        numbered_item(doc, 0, s)

    body_para(doc, "Power_Input format:  C{card}_{type}{num}_{value}")
    p = doc.add_paragraph(style='List Bullet')
    inline_code(p, "C2_AO2_10")
    add_run_normal(p, "  →  Card 2, Analog Output 2, set to 10 V", size_pt=10, color=C_GRAY)
    p = doc.add_paragraph(style='List Bullet')
    inline_code(p, "C3_DO5_1")
    add_run_normal(p, "  →  Card 3, Digital Output 5, set to HIGH", size_pt=10, color=C_GRAY)

    body_para(doc, "If Power_Input is empty, the voltage is measured passively.")

    section_heading(doc, "9.2", "Pull-Up Test", level=2)
    body_para(doc, "Purpose: Verify that a pull-up resistor network is functional.")
    note_box(doc, "Prerequisite: Power test must have passed AND Power_Expected must be 0 V. "
                  "If power is present on the line the pull-up test is automatically skipped.")
    body_para(doc, "Procedure:")
    for s in ["The pull-up pin (from PullUp_Input) is activated via the IO card digital output",
              "The resulting voltage is measured",
              "Measured value is compared to PullUp_Expected ± tolerance"]:
        numbered_item(doc, 0, s)

    section_heading(doc, "9.3", "Logic Test", level=2)
    body_para(doc, "Purpose: Verify discrete digital input/output functionality.")
    body_para(doc, "Procedure:")
    for s in ["A digital command is sent to the IO card (Logic_Command)",
              "The resulting digital state of the pin is read back as a digital input",
              "Compared to Logic_Expected (\"High\" or \"Low\")"]:
        numbered_item(doc, 0, s)

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 10 — I-BIT
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "10.", "I-BIT TEST (BUILT-IN TEST)")

    body_para(doc,
        "The I-BIT test automatically scans all 50 connector pins on both system A "
        "and system B (100 measurements total). A test file does not need to be loaded.")

    body_para(doc, "Purpose:")
    bullet(doc, "Detect short-circuit or open-circuit faults across the full connector before running individual pin tests")

    body_para(doc, "Procedure per pin:")
    for s in ["Mux matrix routes the pin to the measurement ADC",
              "Voltage measured BEFORE pull-up (expected: ~0 V)",
              "Pull-up activated",
              "Voltage measured AFTER pull-up (expected: ~4 V)",
              "Both measurements compared to tolerances from settings"]:
        numbered_item(doc, 0, s)

    body_para(doc, "Results:")
    bullet(doc, "IBIT button turns GREEN — all 100 pins passed both measurements")
    bullet(doc, "IBIT button turns RED   — one or more pins failed")
    bullet(doc, "Full per-pin results are listed in the Operational Log")

    note_box(doc, "Runtime: approximately 3–8 minutes depending on stabilisation delays configured in Comm_settings.yaml.")
    body_para(doc, "To stop early: click  [ Stop IBIT ]")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 11 — Log Window
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "11.", "LOG WINDOW")

    body_para(doc, "The Operational Log records all application events in real time.")
    body_para(doc, "Message format:  [HH:MM:SS]  message text")

    body_para(doc, "Colour coding:")
    log_colours = [
        ("White / grey",  "INFO (INF) — general events"),
        ("Green",         "SUCCESS (SUC) — tests passed"),
        ("Yellow/orange", "WARNING (WRN) — tests failed, pins skipped"),
        ("Red",           "ERROR (ERR) — hardware faults, file errors"),
        ("Dark grey",     "DEBUG (DBG) — verbose internal messages"),
    ]
    for colour, desc in log_colours:
        p = doc.add_paragraph(style='List Bullet')
        add_run_bold(p, colour + ":  ", size_pt=10, color=C_ACCENT)
        add_run_normal(p, desc, size_pt=10, color=C_GRAY)
        set_paragraph_spacing(p, before=0, after=2)

    body_para(doc,
        "Recommended settings during normal operation: INF ☑  SUC ☑  WRN ☑  ERR ☑  DBG ☐")
    body_para(doc,
        "[ Clear ] removes all log text. Test results in the pin table are unaffected.")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 12 — Report
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "12.", "SAVING A REPORT")

    body_para(doc, "Click  [ Report ]  after a test run to save results to an Excel file.")
    body_para(doc, "The report file contains:")
    bullet(doc, "All 22 pin table columns")
    bullet(doc, "All test results (Power, PullUp, Logic) and reason text")
    bullet(doc, "A timestamp row at the top")

    body_para(doc, "Recommended naming convention:")
    p = doc.add_paragraph()
    inline_code(p, "UUT_serial_ConnectorID_YYYY-MM-DD.xlsx")
    add_run_normal(p, "   e.g.  ", size_pt=10, color=C_GRAY)
    inline_code(p, "SN1042_J1_2026-06-25.xlsx")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 13 — Configuration
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "13.", "CONFIGURATION (Comm_settings.yaml)")

    body_para(doc, "Located in:  config\\Comm_settings.yaml  (same folder as the EXE)")
    body_para(doc, "Click  [ Settings ]  to open in a text editor.")

    section_heading(doc, "13.1", "Board Section", level=2)
    body_para(doc, "Controls serial communication with the Controllino.")

    board_rows = [
        ("Type",       "ControllinoMega",  "Hardware type — do not change"),
        ("Port",       "COM5",             "Serial COM port of the Controllino USB cable"),
        ("BaudRate",   "115200",           "Serial baud rate — do not change"),
        ("simulation", "false",            "true = simulation mode, false = real hardware"),
    ]
    tbl = doc.add_table(rows=1 + len(board_rows), cols=3)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Key")
    th_cell(tbl.rows[0].cells[1], "Default")
    th_cell(tbl.rows[0].cells[2], "Description")
    for i, (key, default, desc) in enumerate(board_rows):
        alt = (i % 2 == 1)
        td_cell(tbl.rows[i+1].cells[0], key,     alt=alt, bold=True, size_pt=9)
        td_cell(tbl.rows[i+1].cells[1], default, alt=alt, size_pt=9)
        td_cell(tbl.rows[i+1].cells[2], desc,    alt=alt, size_pt=9)
    doc.add_paragraph()

    note_box(doc,
        "IMPORTANT: Change Port to match the COM port assigned by Windows. "
        "To find it: Device Manager → Ports (COM & LPT)")

    section_heading(doc, "13.2", "UDP_Settings Section", level=2)
    body_para(doc, "Controls communication with IO cards C1–C7.")

    udp_rows = [
        ("Frequency_Hz",         "20.0",  "Communication rate to IO cards"),
        ("Communication_Timeout","2.0",   "Seconds before a card is considered unresponsive"),
        ("localhost_mode",       "false", "true = use loopback addresses (127.0.0.1)"),
    ]
    tbl = doc.add_table(rows=1 + len(udp_rows), cols=3)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Key")
    th_cell(tbl.rows[0].cells[1], "Default")
    th_cell(tbl.rows[0].cells[2], "Description")
    for i, (key, default, desc) in enumerate(udp_rows):
        alt = (i % 2 == 1)
        td_cell(tbl.rows[i+1].cells[0], key,     alt=alt, bold=True, size_pt=9)
        td_cell(tbl.rows[i+1].cells[1], default, alt=alt, size_pt=9)
        td_cell(tbl.rows[i+1].cells[2], desc,    alt=alt, size_pt=9)
    doc.add_paragraph()

    body_para(doc, "Per-card settings (repeated for each of C1–C7):")
    card_rows = [
        ("card_id",      "1",               "Card number"),
        ("enabled",      "true / false",    "Include this card in communication"),
        ("send_ip",      "192.168.195.11",  "IP address of the IO card"),
        ("send_port",    "2880",            "UDP port to send commands"),
        ("receive_ip",   "192.168.195.101", "IP address of this PC"),
        ("receive_port", "1011",            "UDP port to receive responses"),
    ]
    tbl = doc.add_table(rows=1 + len(card_rows), cols=3)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Key")
    th_cell(tbl.rows[0].cells[1], "Example")
    th_cell(tbl.rows[0].cells[2], "Description")
    for i, (key, example, desc) in enumerate(card_rows):
        alt = (i % 2 == 1)
        td_cell(tbl.rows[i+1].cells[0], key,     alt=alt, bold=True, size_pt=9)
        td_cell(tbl.rows[i+1].cells[1], example, alt=alt, size_pt=9)
        td_cell(tbl.rows[i+1].cells[2], desc,    alt=alt, size_pt=9)
    doc.add_paragraph()

    note_box(doc, "After editing Comm_settings.yaml, restart the application for changes to take effect.")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 14 — Simulation and LocalHost
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "14.", "SIMULATION AND LOCALHOST MODES")

    section_heading(doc, "14.1", "Simulation Mode", level=2)
    body_para(doc,
        "Activated by clicking  [ Simulation: On ]  in the Test/Debug group, "
        "or by setting  Board.simulation: true  in Comm_settings.yaml.")
    body_para(doc, "In simulation mode:")
    bullet(doc, "No serial communication with the Controllino occurs")
    bullet(doc, "Voltage measurements return random values within realistic ranges")
    bullet(doc, "Tests complete normally and produce Pass/Fail results")
    bullet(doc, "Use for demonstrations, training, or office testing without hardware")
    body_para(doc, "Button label shows current state:  \"Simulation: On\" (active)  or  \"Simulation: Off\" (real hardware).")

    section_heading(doc, "14.2", "LocalHost Mode", level=2)
    body_para(doc,
        "Activated by clicking  [ LocalHost ]  in the Test/Debug group, "
        "or by setting  UDP_Settings.localhost_mode: true  in Comm_settings.yaml.")
    body_para(doc, "In localhost mode:")
    bullet(doc, "IO card commands are sent to 127.0.0.1 instead of the real card IPs")
    bullet(doc, "A local simulator responds to the commands")
    bullet(doc, "The Controllino (serial) still communicates normally unless Simulation mode is also active")
    bullet(doc, "Use for bench testing when IO cards are not present on the network")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 15 — Debug Mode
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "15.", "DEBUG / STEP MODE")

    body_para(doc,
        "Debug mode allows step-by-step execution of test procedures for verification or fault-finding.")

    body_para(doc, "Activating:  click  [ Debug: False ]  to toggle to  [ Debug: True ].")
    body_para(doc,
        "When active: the test pauses at internal checkpoints. The log shows: "
        "\"in node {N} Waiting for Next button press...\"")
    body_para(doc, "Advancing: click  [ Next ]  to proceed to the next checkpoint.")

    body_para(doc, "HTML Visualization:")
    for s in ["Select an HTML file in the dropdown (Test/Debug group)",
              "Click  [ DOC ]  to open the flow diagram in your browser",
              "The currently active node is highlighted automatically",
              "The browser polls for updates every 500 ms"]:
        numbered_item(doc, 0, s)

    body_para(doc, "Turning off: click  [ Debug: True ]  to toggle back to  [ Debug: False ].")

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 16 — Understanding Results
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "16.", "UNDERSTANDING TEST RESULTS")

    section_heading(doc, "16.1", "Result Values", level=2)
    res_rows = [
        ("Pass", "Measured value is within the defined tolerance of the expected value"),
        ("Fail", "Measured value is outside the tolerance band"),
        ("---",  "Test was not run (column empty in file, or prerequisite not met)"),
    ]
    tbl = doc.add_table(rows=len(res_rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (val, desc) in enumerate(res_rows):
        td_cell(tbl.rows[i].cells[0], val, alt=(i%2==1), bold=True, size_pt=10)
        td_cell(tbl.rows[i].cells[1], desc, alt=(i%2==1), size_pt=10)
    doc.add_paragraph()

    section_heading(doc, "16.2", "Pull-Up Test Skip Conditions", level=2)
    body_para(doc, "The PullUp test is automatically skipped if:")
    bullet(doc, "Power_Expected is not 0 V (the pin has DC power on it)")
    bullet(doc, "The Power test for that pin failed")
    body_para(doc, "The Operational Log will explain the skip with a WARNING message.")

    section_heading(doc, "16.3", "Voltage Tolerance", level=2)
    body_para(doc, "The tolerance is configured in Comm_settings.yaml:  Test.voltage_degredation  (default: 3.0 V)")
    body_para(doc, "A measurement is Pass if:  |measured − expected| ≤ tolerance")

    section_heading(doc, "16.4", "Common Failure Reasons", level=2)
    failures = [
        ("Measurement is out of tolerance",
         "Voltage differs from expected by more than the configured tolerance. Check wiring and UUT."),
        ("Analog pin not found in pin map",
         "The board configuration does not have a mapping for the required channel. Contact support."),
        ("Failed to set mux bits",
         "The Controllino did not respond correctly to the matrix command. Check connection (Comm Check)."),
        ("Relay did not operate",
         "(I-BIT only) The general relay did not close/open as expected. Check relay wiring."),
    ]
    tbl = doc.add_table(rows=1 + len(failures), cols=2)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Reason Text")
    th_cell(tbl.rows[0].cells[1], "Explanation & Action")
    for i, (reason, action) in enumerate(failures):
        td_cell(tbl.rows[i+1].cells[0], reason, alt=(i%2==1), bold=True, size_pt=9)
        td_cell(tbl.rows[i+1].cells[1], action, alt=(i%2==1), size_pt=9)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 17 — Troubleshooting
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "17.", "TROUBLESHOOTING")

    ts_rows = [
        ("Hardware warning at startup",
         "Wrong COM port or USB cable not seated",
         "Open Settings, change Board.Port to correct COM; reconnect cable"),
        ("Comm Check fails",
         "Controllino not ready or firmware issue",
         "Unplug/replug USB, restart Controllino"),
        ("UDP binding error on startup",
         "Port already in use by another process",
         "Close other IO Tester instances; reboot PC"),
        ("All pins Fail",
         "Simulation mode ON / wrong connector map / UUT not seated",
         "Toggle Simulation Off; verify file; check fixture seating"),
        ("Test / Test_All greyed out",
         "No file loaded",
         "Click Load first"),
        ("Pull-up tests all skipped",
         "Power_Expected ≠ 0 V in test file",
         "Pull-up only runs when Power Expected = 0 V for that pin"),
        ("Application does not start",
         "Missing config files",
         "Verify config\\ folder is present next to the EXE"),
        ("Log window is empty",
         "All log filters unchecked",
         "Check INF/SUC/WRN/ERR checkboxes in the Log group"),
    ]

    tbl = doc.add_table(rows=1 + len(ts_rows), cols=3)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Symptom")
    th_cell(tbl.rows[0].cells[1], "Likely Cause")
    th_cell(tbl.rows[0].cells[2], "Action")
    for i, (symptom, cause, action) in enumerate(ts_rows):
        alt = (i % 2 == 1)
        td_cell(tbl.rows[i+1].cells[0], symptom, alt=alt, bold=True, size_pt=9)
        td_cell(tbl.rows[i+1].cells[1], cause,   alt=alt, size_pt=9)
        td_cell(tbl.rows[i+1].cells[2], action,  alt=alt, size_pt=9)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # Section 18 — Shortcuts
    # ═══════════════════════════════════════════════════════════════════════════
    section_heading(doc, "18.", "KEYBOARD AND MOUSE SHORTCUTS")

    shortcuts = [
        ("Click row",              "Select single row for testing"),
        ("Ctrl + Click",           "Add or remove a row from the selection"),
        ("Shift + Click",          "Select a range of rows"),
        ("Click column header",    "Sort by that column (click again to reverse)"),
        ("Double-click cell",      "Open inline editor (editable columns only)"),
        ("Hover over cell",        "Show tooltip with full cell text"),
    ]
    tbl = doc.add_table(rows=1 + len(shortcuts), cols=2)
    tbl.style = 'Table Grid'
    th_cell(tbl.rows[0].cells[0], "Action")
    th_cell(tbl.rows[0].cells[1], "Result")
    for i, (action, result) in enumerate(shortcuts):
        td_cell(tbl.rows[i+1].cells[0], action, alt=(i%2==1), bold=True, size_pt=10)
        td_cell(tbl.rows[i+1].cells[1], result, alt=(i%2==1), size_pt=10)
    doc.add_paragraph()

    note_box(doc, "There are no keyboard shortcuts for toolbar buttons — use the mouse.")

    # ── Footer ────────────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_normal(p, "IO Tester  |  User Manual  |  Version 1.0  |  June 2026",
                   size_pt=9, color=C_GRAY, italic=True)

    doc.save(r"c:\ArduinoProject\IO_Tester\docs\User_Manual.docx")
    print("User_Manual.docx saved.")


# ─────────────────────────────────────────────
if __name__ == "__main__":
    build_quick_start()
    build_user_manual()
    print("Done.")
